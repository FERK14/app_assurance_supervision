
import streamlit as st
import pandas as pd
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import StandardScaler
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from openai import OpenAI

# Configurer OpenAI (version >= 1.0.0)
client = OpenAI(api_key=st.secrets["openai_api_key"])

st.set_page_config(page_title="Supervision IA - Assurances", layout="wide")

# Charger les données
df = pd.read_excel("donnees_assurances_simulees.xlsx")

# Préparer les données
X = df.drop(columns=["Nom_Compagnie", "En_risque"])
scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

# Entraîner le modèle
model = RandomForestClassifier(n_estimators=100, random_state=42)
model.fit(X_scaled, df["En_risque"])
df["Prediction_IA"] = model.predict(X_scaled)
df["Probabilité_Risque"] = model.predict_proba(X_scaled)[:, 1]

# Générer commentaire IA avec OpenAI GPT
def rediger_commentaire(row):
    prompt = f"""
Tu es analyste du secteur des assurances. Rédige un commentaire professionnel de 150 mots sur cette compagnie :

- Chiffre d'affaires : {row['Chiffre_affaires']:,} FCFA
- Taux de sinistralité : {row['Taux_sinistralite']:.2f} %
- Ratio de solvabilité : {row['Ratio_solvabilite']:.2f} %
- Ratio de liquidité : {row['Ratio_liquidite']:.2f} %
- Nombre d'agences : {int(row['Nombre_agences'])}
- Rendement des investissements : {row['Rendement_investissements']:.2f} %
- Score IA (risque) : {row['Probabilité_Risque']:.2%}

Analyse les risques, les points forts, et propose une recommandation claire pour les autorités de régulation.
"""
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Erreur OpenAI : {e}"

# Interface utilisateur
st.title("📊 Supervision IA du secteur des assurances")

# SECTION 1 : Analyse individuelle
st.header("🔍 Analyse d'une compagnie")
index = st.selectbox("Sélectionner une compagnie", df.index)
row = df.loc[index]

st.subheader("📌 Indicateurs financiers")
cols1, cols2, cols3 = st.columns(3)
cols1.metric("Chiffre d'affaires", f"{row['Chiffre_affaires']:,} FCFA")
cols2.metric("Sinistralité (%)", f"{row['Taux_sinistralite']:.2f}")
cols3.metric("Solvabilité (%)", f"{row['Ratio_solvabilite']:.2f}")
cols4, cols5, cols6 = st.columns(3)
cols4.metric("Liquidité (%)", f"{row['Ratio_liquidite']:.2f}")
cols5.metric("Agences", int(row["Nombre_agences"]))
cols6.metric("Rendement Invest. (%)", f"{row['Rendement_investissements']:.2f}")

st.subheader("🧠 Prédiction IA")
if row["Prediction_IA"] == 1:
    st.error(f"⚠️ Risque détecté (probabilité : {row['Probabilité_Risque']:.2%})")
else:
    st.success(f"✅ Aucun risque détecté (probabilité : {row['Probabilité_Risque']:.2%})")

st.subheader("📈 Ratio de solvabilité (comparaison)")
fig, ax = plt.subplots(figsize=(10, 6))
sns.barplot(x=df["Ratio_solvabilite"], y=df["Nom_Compagnie"], palette="coolwarm", ax=ax)
ax.axvline(100, color="red", linestyle="--")
ax.set_xlabel("Ratio de solvabilité (%)")
st.pyplot(fig)

# SECTION 2 : Rapport global Word
st.header("📝 Rapport IA - Top 10 compagnies à risque")

if st.button("📄 Générer le rapport Word"):
    top10 = df.sort_values(by="Probabilité_Risque", ascending=False).head(10)
    doc = Document()
    doc.add_heading("Rapport IA – Top 10 compagnies à risque", 0)

    with st.spinner("⏳ Génération des commentaires IA..."):
        for i, row in top10.iterrows():
            commentaire = rediger_commentaire(row)
            doc.add_heading(row["Nom_Compagnie"], level=1)
            doc.add_paragraph(f"Chiffre d'affaires : {row['Chiffre_affaires']:,} FCFA")
            doc.add_paragraph(f"Score IA : {row['Probabilité_Risque']:.2%}")
            doc.add_paragraph("📝 Commentaire IA :")
            doc.add_paragraph(commentaire)

    filename = "rapport_assurances_ia.docx"
    doc.save(filename)
    with open(filename, "rb") as f:
        st.download_button("📥 Télécharger le rapport", f, file_name=filename)
